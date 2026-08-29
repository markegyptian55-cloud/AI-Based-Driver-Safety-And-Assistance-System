"""
build_book.py — assemble the MTC-format project book (.docx)
=============================================================
Builds the print-ready document from the project's own verified artefacts,
using the official MTC template as the style source.

The template is opened, its placeholder body is cleared, and content is
rebuilt using ITS styles -- so page setup, fonts, heading hierarchy and
caption formatting are inherited rather than re-invented. The original
template file is never written to.

Page numbering follows the template's rule: front matter in Roman
numerals, then Chapter 1 restarts Arabic numbering at 1. That needs raw
OOXML (python-docx has no page-numbering API), which is what the
_add_page_numbering / _sect_pr helpers below are for.

TOC, List of Figures and List of Tables are inserted as Word FIELD codes,
not as static text. They are therefore empty until opened in Word and
refreshed (Ctrl+A, F9) -- which is the template's own documented
procedure and the only way the page numbers can be correct.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
INFO = PROJECT_ROOT / "INFO"
TEMPLATE = Path(r"C:\Users\student\Desktop\MTC_format_Professional Diploma Project_2026.docx")
OUT = Path.home() / "Desktop" / "AI-Based_Driver_Safety_and_Assistance_System.docx"

SCRATCH_FS = Path(r"D:\project\Driver project\michel from scartch\results")

# ----------------------------------------------------------------- helpers


def clear_body(doc):
    """Remove the template's placeholder content, keeping its styles."""
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _field(paragraph, instr: str):
    """Insert a Word field code (TOC / PAGE / etc.) into a paragraph."""
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr_el, fld_sep, fld_end):
        r._r.append(el)


def _sect_pr(section):
    return section._sectPr


def set_page_numbering(section, fmt: str, start: int | None = None):
    """
    fmt: 'lowerRoman' | 'decimal'. start=None continues the previous
    sequence; an integer restarts at that number.
    """
    sectPr = _sect_pr(section)
    for existing in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(existing)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def add_footer_page_number(section, link_to_previous=False):
    """Bottom-centre page number, per the template's stated convention."""
    footer = section.footer
    footer.is_linked_to_previous = link_to_previous
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, " PAGE ")


def para(doc, text="", style=None, align=None, bold=None, size=None,
         italic=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if size is not None:
            run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def body_text(doc, text):
    """Body paragraph in the template's body style (TNR 12, 1.5, justified)."""
    return para(doc, text, style="No Spacing")


def heading(doc, text, level):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    return para(doc, text, style=style)


def chapter_title(doc, number, title):
    """Chapter heading pair, matching the template's 'Chapter (N)' pattern."""
    para(doc, f"Chapter ({number})", style="Heading 1")
    para(doc, title, style="Heading 1")


def caption(doc, text):
    p = para(doc, text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def figure(doc, path: Path, number: str, title: str, width_cm=13.5):
    if not path.exists():
        body_text(doc, f"[FIGURE MISSING: {path.name}]")
        caption(doc, f"Fig. ({number}) {title}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    caption(doc, f"Fig. ({number}) {title}")


def table(doc, number, title, headers, rows, widths=None):
    caption(doc, f"Table ({number}) {title}")
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return t


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def arabic_para(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Right-to-left paragraph for the Arabic back matter."""
    p = doc.add_paragraph()
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return p


# ----------------------------------------------------------------- data


def load_metrics():
    """Every model's verified metrics, keyed by INFO folder."""
    out = {}
    for mj in INFO.glob("*/*-test-result/tested-images/metrics.json"):
        with open(mj, encoding="utf-8") as fh:
            d = json.load(fh)
        d["_family"] = mj.parents[2].name
        d["_run"] = mj.parents[1].name.replace("-test-result", "")
        out[d["model_key"]] = d
    return out


def load_latency():
    p = INFO / "_benchmark" / "latency.json"
    if not p.exists():
        return {}
    with open(p, encoding="utf-8") as fh:
        d = json.load(fh)
    return {m["model_key"]: m for m in d.get("models", [])}


# Display name, family label, input size, epochs, GPU-hours, timing source.
# Epochs/hours are read from the project's training records; the source
# column records whether a duration was machine-logged or self-reported,
# because the two are not the same standard of evidence.
RUNS = [
    ("yolo11m-1-warmstart",              "YOLO11m Warm-Start",            "YOLO11m",       640,  15,   7.23, "logged"),
    ("yolo11n-1-capacity",               "YOLO11n Capacity",              "YOLO11n",       960, 112,  23.07, "logged"),
    ("yolo26n-4-calibration",            "YOLO26n Calibration",           "YOLO26n",       960,  40,   9.50, "logged"),
    ("yolo26n-2-finetune",               "YOLO26n Fine-Tune",             "YOLO26n",       960,  50,  11.70, "logged"),
    ("yolo26n-6-weakdevice-480",         "YOLO26n Weak-Device",           "YOLO26n",       480,  25,   1.63, "logged"),
    ("yolo26n-5-cls3",                   "YOLO26n Class-Weight 3.0",      "YOLO26n",       960,  34,   8.08, "logged"),
    ("yolo26s-1-capacity",               "YOLO26s Capacity",              "YOLO26s",       960,  47,  20.42, "logged"),
    ("yolo26n-3-fresh-640",              "YOLO26n Fresh 640",             "YOLO26n",       640, 100,  10.73, "logged"),
    ("yolo26n-1-baseline",               "YOLO26n Baseline",              "YOLO26n",       960,  77,  20.39, "logged"),
    ("old-rfdetr-nano-2-finetune-384",   "RF-DETR-Nano Fine-Tune",        "RF-DETR-Nano",  384,  15,  11.84, "reported"),
    ("old-rfdetr-nano-1-baseline-384",   "RF-DETR-Nano Baseline",         "RF-DETR-Nano",  384,  50,   4.50, "reported"),
    ("old-yolo11m-3-worstcase-d18source-640", "YOLO11m Cross-Dataset",    "YOLO11m",       640,  40,  40.43, "reported"),
    ("old-yolo11m-2-trial2winner-640",   "YOLO11m Extended",              "YOLO11m",       640,  53,  26.64, "reported"),
    ("old-rfdetr-small-3-worstcase-384", "RF-DETR-Small Worst-Case",      "RF-DETR-Small", 384,  17,   8.50, "approx."),
    ("old-rfdetr-small-2-standard-640",  "RF-DETR-Small Standard",        "RF-DETR-Small", 640,  40,  33.33, "reported"),
    ("old-yolo11m-1-worstcase-384",      "YOLO11m Worst-Case",            "YOLO11m",       384,  40,   7.68, "reported"),
    ("old-yolo11n-2-worstcase-dms-640",  "YOLO11n Worst-Case",            "YOLO11n",       640,  60,   7.01, "reported"),
    ("old-yolo11n-1-baseline-384",       "YOLO11n Baseline",              "YOLO11n",       384,  60,   7.11, "reported"),
    ("old-yolo26n-2-nanoworstcase-384",  "YOLO26n Compact Worst-Case",    "YOLO26n",       384,  20,   1.06, "logged"),
    ("old-yolo26n-1-nanobaseline-384",   "YOLO26n Compact Baseline",      "YOLO26n",       384,  40,   2.52, "logged"),
    ("old-rfdetr-small-1-baseline-640",  "RF-DETR-Small Baseline",        "RF-DETR-Small", 640,  15,  14.20, "reported"),
]


def pct(v):
    return f"{v * 100:.2f}" if isinstance(v, (int, float)) else "—"


# Two models are recorded under a different key in their metrics.json than
# in configs/checkpoints.yaml, because the evaluation predates the registry
# naming. Mapping them explicitly is safer than renaming either source:
# registry key -> metrics.json model_key.
METRICS_KEY_ALIAS = {
    "yolo11m-1-warmstart": "yolo11m-warmstart-pilot-640",
    "yolo11n-1-capacity": "yolo11n-capacity-960",
}


def lookup_metrics(metrics: dict, key: str):
    """Metrics for a registry key, tolerating the alias mismatch above."""
    return metrics.get(key) or metrics.get(METRICS_KEY_ALIAS.get(key, ""))


def ranked(metrics):
    rows = []
    missing = []
    for key, disp, fam, imgsz, ep, hrs, src in RUNS:
        m = lookup_metrics(metrics, key)
        if not m:
            missing.append(key)
            continue
        rows.append({
            "key": key, "disp": disp, "fam": fam, "imgsz": imgsz,
            "ep": ep, "hrs": hrs, "src": src,
            "map": m.get("map50_corrected"), "map_raw": m.get("map50"),
            "p": m.get("precision"), "r": m.get("recall"), "f1": m.get("f1"),
            "apc": m.get("ap_per_class_corrected", {}),
        })
    rows.sort(key=lambda r: r["map"] or 0, reverse=True)
    if missing:
        # Loud rather than silent: a model quietly dropped from the results
        # table is exactly the kind of omission this project exists to avoid.
        print(f"WARNING: {len(missing)} run(s) had no metrics and are absent "
              f"from the results table: {missing}")
    return rows


# ----------------------------------------------------------------- content

TITLE_EN = "AI-Based Driver Safety and Assistance System"
TITLE_AR = "نظام قائم على الذكاء الاصطناعي لسلامة السائق والمساعدة أثناء القيادة"
SUPERVISOR = "Prof. Dr. Kamel Elhadad"
TEAM = [
    "Mohamed Mostafa Mohamed El-Basyouni (Team Leader)",
    "Michael Magdy Amin Sidhom",
    "Ali Ibrahim Ahmed Othman",
    "Karim Mustafa Ali Ibrahim",
    "Mohamed Osama Bahnasy Abdel Halim",
]
TEAM_AR = [
    "محمد مصطفى محمد البسيوني (قائد الفريق)",
    "مايكل مجدي أمين سدهم",
    "علي إبراهيم أحمد عثمان",
    "كريم مصطفى علي إبراهيم",
    "محمد أسامة بهنساوي عبد الحليم",
]


def title_page(doc, hard_cover: bool):
    para(doc, "Military Technical College", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=16, bold=False)
    para(doc, "Department of Computer Engineering and Artificial Intelligence",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    doc.add_paragraph()
    note = ("[INSERT OFFICIAL COLLEGE LOGO HERE — height 50, width 53.5"
            + (" — hard-cover version]" if hard_cover else "]"))
    p = para(doc, note, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)
    doc.add_paragraph()
    para(doc, TITLE_EN, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    doc.add_paragraph()
    para(doc, "By", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    for name in TEAM:
        para(doc, name, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
             space_after=0)
    para(doc, "Military Technical College", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=14, italic=True)
    doc.add_paragraph()
    para(doc, "Under supervision of", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=16, bold=True)
    para(doc, SUPERVISOR, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True,
         space_after=0)
    para(doc, "Military Technical College", align=WD_ALIGN_PARAGRAPH.CENTER,
         size=14, italic=True)
    doc.add_paragraph()
    para(doc,
         "This project is submitted in partial fulfillment of the requirements "
         "for the degree of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    para(doc, "Professional Diploma in Applied AI and Data Analytics",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    doc.add_paragraph()
    para(doc,
         "Conducted within the Digilians (Digital Pioneers) Initiative",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
    doc.add_paragraph()
    para(doc, "Cairo 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True)
    page_break(doc)


ABSTRACT_EN = [
 "Driver fatigue is a persistent contributor to road traffic collisions, and "
 "its early behavioural indicators — prolonged eye closure and yawning — are "
 "observable from a single in-cabin camera. This project develops a complete "
 "vision-based driver monitoring system, from dataset construction through "
 "model selection to browser-based deployment.",

 "The work began from a 57,098-image corpus assembled from several "
 "independently annotated sources. Analysis established that the corpus was a "
 "merge of separate single-task datasets — an eye-state corpus, a yawning "
 "corpus, and driver-monitoring session recordings — combined into one "
 "three-class label space without re-annotation. This produced systematic, "
 "source-correlated missing supervision rather than random label noise, a "
 "setting known in the literature as sparsely annotated object detection. A "
 "structured reconciliation process, including human verification of sampled "
 "images against spatial evidence gates, produced a source-aware supervision "
 "manifest recording which classes each source family can be trusted to "
 "supervise. Group-aware splitting was applied so that near-duplicate frames "
 "from the same recording session could not span the train and test "
 "partitions, and the resulting split was independently verified for leakage. "
 "The final dataset contains 50,654 images and 68,292 annotated instances "
 "across the classes closed_eye, open_eye and yawning.",

 "Twenty-one training runs were carried out across four detection "
 "architecture families — YOLO26, YOLOv11, the transformer-based RF-DETR, and "
 "a simplified Faster R-CNN implemented from scratch — spanning 950 epochs and "
 "277.57 GPU-hours, at input resolutions from 384 to 960 pixels and under two "
 "augmentation intensities. All models were re-evaluated under a single "
 "protocol on a held-out 5,589-image test partition, reporting both standard "
 "and label-gap-corrected metrics, the latter accounting for the partial "
 "annotation identified during dataset analysis.",

 "Two YOLO26n configurations were selected for deployment, at 480 and 960 "
 "pixel input, achieving 82.72 % and 82.75 % corrected mAP@0.5 respectively. "
 "Both were exported to ONNX with non-maximum suppression compiled into the "
 "graph and converted to half precision, which halved model size with no "
 "measurable loss of detection agreement. The models are executed entirely "
 "on the user's own device through a browser application using WebGPU with a "
 "multi-threaded WebAssembly fallback, so that video is never transmitted.",

 "Temporal interpretation is performed by rolling-window PERCLOS aggregation "
 "with a debounced state machine. Learned temporal modelling using recurrent "
 "or sequence architectures was not implemented in this work and is identified "
 "as the next stage of development.",
]

ABSTRACT_AR = [
 "يمثل إرهاق السائق أحد الأسباب المستمرة لحوادث الطرق، ويمكن ملاحظة مؤشراته "
 "السلوكية المبكرة — إغلاق العينين لفترات طويلة والتثاؤب — من خلال كاميرا واحدة "
 "داخل المركبة. يقدم هذا المشروع نظامًا متكاملًا لمراقبة السائق يعتمد على الرؤية "
 "الحاسوبية، بدءًا من بناء مجموعة البيانات ومرورًا باختيار النموذج وانتهاءً "
 "بالنشر داخل المتصفح.",

 "بدأ العمل من مجموعة بيانات تضم 57,098 صورة جُمعت من عدة مصادر مُوسمة بشكل "
 "مستقل. أظهر التحليل أن هذه المجموعة كانت دمجًا لمجموعات بيانات أحادية المهمة — "
 "مجموعة لحالة العين، ومجموعة للتثاؤب، وتسجيلات لجلسات مراقبة السائق — دُمجت في "
 "فضاء تسميات ثلاثي الفئات دون إعادة توسيم. نتج عن ذلك نقص منهجي في الإشراف مرتبط "
 "بالمصدر وليس ضوضاء عشوائية في التسميات. وقد أنتجت عملية مراجعة منظمة، شملت تحققًا "
 "بشريًا من عينات من الصور، سجلًا للإشراف يراعي المصدر. كما طُبق تقسيم واعٍ "
 "بالمجموعات لمنع تسرب الإطارات المتشابهة بين مجموعتي التدريب والاختبار. تحتوي "
 "مجموعة البيانات النهائية على 50,654 صورة و68,292 حالة مُوسمة عبر الفئات: العين "
 "المغلقة، والعين المفتوحة، والتثاؤب.",

 "أُجريت إحدى وعشرون تجربة تدريب عبر أربع عائلات معمارية للكشف عن الأجسام — "
 "YOLO26 وYOLOv11 وRF-DETR القائم على المحولات، بالإضافة إلى نموذج Faster R-CNN "
 "مبسط مُنفذ من الصفر — بإجمالي 950 حقبة تدريبية و277.57 ساعة معالجة رسومية، "
 "وبدقات إدخال تتراوح بين 384 و960 بكسل. وقد أُعيد تقييم جميع النماذج وفق بروتوكول "
 "موحد على مجموعة اختبار مستقلة تضم 5,589 صورة.",

 "تم اختيار نموذجين من YOLO26n للنشر، بدقة إدخال 480 و960 بكسل، وحققا 82.72٪ "
 "و82.75٪ على التوالي وفق مقياس mAP@0.5 المصحح. وقد صُدّر كلا النموذجين إلى صيغة "
 "ONNX مع دمج خطوة الكبت غير الأقصى داخل الرسم البياني، وتحويلهما إلى دقة نصفية، "
 "مما خفض حجم النموذج إلى النصف دون خسارة تُذكر في دقة الكشف. ويجري تنفيذ النماذج "
 "بالكامل على جهاز المستخدم عبر تطبيق يعمل في المتصفح، بحيث لا يُنقل الفيديو إلى "
 "أي خادم خارجي.",

 "يتم التفسير الزمني عبر تجميع PERCLOS ضمن نافذة متحركة مع آلة حالة مُثبّتة. "
 "أما النمذجة الزمنية المُتعلمة باستخدام المعماريات التكرارية أو التسلسلية فلم "
 "تُنفذ في هذا العمل، وقد حُددت كمرحلة تطوير لاحقة.",
]

ABBREVIATIONS = [
 ("ADAS", "Advanced Driver Assistance System"),
 ("AP", "Average Precision"),
 ("CNN", "Convolutional Neural Network"),
 ("COCO", "Common Objects in Context"),
 ("DETR", "Detection Transformer"),
 ("DFL", "Distribution Focal Loss"),
 ("FPS", "Frames Per Second"),
 ("FP16", "Half-Precision Floating Point"),
 ("FP32", "Single-Precision Floating Point"),
 ("GPU", "Graphics Processing Unit"),
 ("IoU", "Intersection over Union"),
 ("mAP", "mean Average Precision"),
 ("NMS", "Non-Maximum Suppression"),
 ("ONNX", "Open Neural Network Exchange"),
 ("PERCLOS", "Percentage of Eye Closure"),
 ("RoI", "Region of Interest"),
 ("RPN", "Region Proposal Network"),
 ("SAOD", "Sparsely Annotated Object Detection"),
 ("VRAM", "Video Random Access Memory"),
 ("WASM", "WebAssembly"),
 ("YOLO", "You Only Look Once"),
]


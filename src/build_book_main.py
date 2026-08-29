"""
build_book_main.py — document assembly
=======================================
Run this to produce the book. Imports helpers from build_book.py and prose
from book_content*.py, then writes the .docx.

    python src/build_book_main.py
"""

from __future__ import annotations

import sys
from pathlib import Path

import docx
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))

from build_book import (  # noqa: E402
    INFO, OUT, PROJECT_ROOT, SCRATCH_FS, TEMPLATE,
    add_footer_page_number, body_text, caption, chapter_title, clear_body,
    figure, heading, load_latency, load_metrics, page_break, para, ranked,
    set_page_numbering, table, title_page, arabic_para, _field, pct,
    TITLE_AR, TITLE_EN, SUPERVISOR, TEAM, TEAM_AR,
)
import book_content as C  # noqa: E402
import book_content2 as C2  # noqa: E402
import book_references as R  # noqa: E402

DEPLOY_480 = "yolo26n-6-weakdevice-480"
DEPLOY_960 = "yolo26n-4-calibration"


def ms(row, dev, field="median_ms"):
    d = (row or {}).get(f"torch_{dev}")
    if not isinstance(d, dict) or "error" in d:
        return "-"
    v = d.get(field)
    return f"{v:.2f}" if isinstance(v, (int, float)) else "-"


def build():
    if not TEMPLATE.exists():
        print(f"Template not found: {TEMPLATE}")
        return 1

    doc = docx.Document(str(TEMPLATE))
    clear_body(doc)

    metrics = load_metrics()
    latency = load_latency()
    rows = ranked(metrics)
    if not rows:
        print("No metrics found under INFO/ - nothing to build.")
        return 1

    # ============================== FRONT MATTER (roman numerals) =========
    sec0 = doc.sections[0]
    set_page_numbering(sec0, "lowerRoman", start=1)
    add_footer_page_number(sec0)

    title_page(doc, hard_cover=True)
    title_page(doc, hard_cover=False)

    # --- Abstract
    para(doc, "Abstract", style="Heading 1 Title")
    for p in C.ABSTRACT_EN:
        body_text(doc, p)
    page_break(doc)

    # --- Acknowledgments
    para(doc, "Acknowledgments", style="Title")
    for p in C.ACKNOWLEDGMENTS:
        body_text(doc, p)
    page_break(doc)

    # --- TOC / LoF / LoT as live Word fields
    para(doc, "Table of Contents", style="Title")
    _field(doc.add_paragraph(), r' TOC \o "1-3" \h \z \u ')
    page_break(doc)

    para(doc, "List of Figures", style="Heading 1 Title")
    _field(doc.add_paragraph(), r' TOC \h \z \c "Figure" ')
    page_break(doc)

    para(doc, "List of Tables", style="Heading 1 Title")
    _field(doc.add_paragraph(), r' TOC \h \z \c "Table" ')
    page_break(doc)

    # --- Nomenclature
    para(doc, "Nomenclature", style="Heading 1 Title")
    body_text(doc, "Abbreviations")
    for abbr, meaning in C.ABBREVIATIONS:
        p = para(doc, f"{abbr}\t{meaning}", style="No Spacing")
        p.paragraph_format.space_after = 0

    # ============================== BODY (arabic, restart at 1) ===========
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_numbering(body_sec, "decimal", start=1)
    add_footer_page_number(body_sec)

    # ---------------- Chapter 1
    chapter_title(doc, 1, "Introduction")
    for h, paras in C.CH1.items():
        heading(doc, h, 2)
        for t in paras:
            body_text(doc, t)
    page_break(doc)

    # ---------------- Chapter 2
    chapter_title(doc, 2, "Literature Review")
    for h, paras in C2.CH2.items():
        heading(doc, h, 2)
        for t in paras:
            body_text(doc, t)
    page_break(doc)

    # ---------------- Chapter 3
    ch3 = INFO / "_book" / "Chapter_3_Methodology.md"
    chapter_title(doc, 3, "Methodology")
    write_chapter3(doc, metrics)
    page_break(doc)

    # ---------------- Chapter 4
    chapter_title(doc, 4, "Implementation and Results")
    write_chapter4(doc, rows, metrics, latency)
    page_break(doc)

    # ---------------- Chapter 5
    chapter_title(doc, 5, "Conclusion")
    for h, paras in C.CH5.items():
        heading(doc, h, 2)
        for t in paras:
            body_text(doc, t)
    page_break(doc)

    # ---------------- References
    para(doc, "References", style="Heading 1")
    for i, ref in enumerate(R.active_references(), 1):
        p = para(doc, f"[{i}]\t{ref}", style="No Spacing")
        p.paragraph_format.space_after = 6
    page_break(doc)

    # ---------------- Appendix
    para(doc, "Appendix (A)", style="Heading 1")
    heading(doc, "A.1 Complete Experimental Record", 2)
    body_text(doc,
              "The complete per-model evidence base - evaluation reports, "
              "metrics, ten diagnostic charts per model, and training curves "
              "where preserved - is retained in the project's INFO directory, "
              "organised by architecture family and run. Table (4.2) "
              "summarises all twenty-one runs.")
    heading(doc, "A.2 Reproducibility", 2)
    body_text(doc,
              "Training configurations for every run are retained as "
              "version-controlled YAML files. The evaluation, benchmarking "
              "and chart-generation procedures are implemented as scripts in "
              "the project's src directory and were used to produce every "
              "figure and table in this document.")
    page_break(doc)

    # ============================== ARABIC BACK MATTER ====================
    ar_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer_page_number(ar_sec, link_to_previous=True)

    arabic_para(doc, "ملخص المشروع", size=16, bold=True)
    doc.add_paragraph()
    for p in C.ABSTRACT_AR:
        arabic_para(doc, p, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    page_break(doc)

    arabic_para(doc, "الكلية الفنية العسكرية", size=16, bold=True)
    arabic_para(doc, "قسم هندسة الحواسب والذكاء الاصطناعي", size=16)
    doc.add_paragraph()
    arabic_para(doc, TITLE_AR, size=18, bold=True)
    doc.add_paragraph()
    arabic_para(doc, "مشروع مقدم من", size=14)
    for n in TEAM_AR:
        arabic_para(doc, n, size=14, bold=True)
    doc.add_paragraph()
    arabic_para(doc, "للحصول على الدبلوم المهني في", size=14)
    arabic_para(doc, "تحليل البيانات والذكاء الاصطناعي التطبيقي", size=14, bold=True)
    doc.add_paragraph()
    arabic_para(doc, "تحت إشراف", size=14)
    arabic_para(doc, "أ.د. كامل الحداد", size=14, bold=True)
    doc.add_paragraph()
    arabic_para(doc, "ضمن مبادرة رواد الرقمية (Digilians)", size=12)
    doc.add_paragraph()
    arabic_para(doc, "القاهرة 2026", size=12, bold=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}")
    print("\nIn Word: Ctrl+A then F9 to populate the Table of Contents,")
    print("List of Figures and List of Tables.")
    return 0


# --------------------------------------------------------------- chapters


def write_chapter3(doc, metrics):
    from book_ch3 import SECTIONS, TABLES
    for key, blocks in SECTIONS.items():
        level = 3 if key.count(".") == 2 else 2
        heading(doc, key, level)
        for b in blocks:
            if isinstance(b, tuple) and b[0] == "TABLE":
                spec = TABLES[b[1]]
                table(doc, spec["n"], spec["title"], spec["headers"], spec["rows"])
            elif isinstance(b, tuple) and b[0] == "FIG":
                figure(doc, b[1], b[2], b[3])
            else:
                body_text(doc, b)


def write_chapter4(doc, rows, metrics, latency):
    heading(doc, "4.1 Experimental Setup and Environment", 2)
    for t in C2.CH4_INTRO:
        body_text(doc, t)
    for t in C2.CH4_ENV:
        body_text(doc, t)

    heading(doc, "4.2 Training Programme and Computational Expenditure", 2)
    for t in C2.CH4_TRAINING:
        body_text(doc, t)

    fam = {}
    for r in rows:
        f = fam.setdefault(r["fam"], {"runs": 0, "ep": 0, "hrs": 0.0, "best": 0})
        f["runs"] += 1
        f["ep"] += r["ep"]
        f["hrs"] += r["hrs"]
        f["best"] = max(f["best"], r["map"] or 0)
    trows = [[k, v["runs"], v["ep"], f"{v['hrs']:.2f}", pct(v["best"])]
             for k, v in sorted(fam.items(), key=lambda x: -x[1]["best"])]
    trows.append(["Total", sum(v["runs"] for v in fam.values()),
                  sum(v["ep"] for v in fam.values()),
                  f"{sum(v['hrs'] for v in fam.values()):.2f}", "-"])
    table(doc, "4.1", "Training effort by architecture family",
          ["Family", "Runs", "Epochs", "GPU-hours", "Best mAP@0.5 (%)"], trows)

    heading(doc, "4.3 Comparative Accuracy", 2)
    for t in C2.CH4_ACCURACY:
        body_text(doc, t)

    arows = []
    for i, r in enumerate(rows, 1):
        arows.append([i, r["disp"], r["fam"], r["imgsz"], pct(r["map_raw"]),
                      pct(r["map"]), pct(r["p"]), pct(r["r"]), pct(r["f1"]),
                      r["ep"], f"{r['hrs']:.2f}", r["src"]])
    table(doc, "4.2", "All training runs, ranked by corrected mAP@0.5",
          ["#", "Model", "Family", "Input", "mAP raw (%)", "mAP corr. (%)",
           "P (%)", "R (%)", "F1 (%)", "Epochs", "GPU-h", "Timing"], arows)
    body_text(doc,
              "Timing column: 'logged' indicates machine-recorded duration; "
              "'reported' indicates a figure self-reported in a run summary; "
              "'approx.' indicates a prose approximation. These are not "
              "aggregated as a single unqualified total.")

    figure(doc, INFO / "_comparison" / "runs_map_comparison.png", "4.1",
           "Comparative mAP@0.5 across all evaluated runs")
    figure(doc, INFO / "_comparison" / "eval_per_class_ap.png", "4.2",
           "Per-class average precision across evaluated runs")

    for t in C2.CH4_ACCURACY_DISCUSS:
        body_text(doc, t)

    heading(doc, "4.4 The Deployed Configurations", 2)
    m480, m960 = metrics.get(DEPLOY_480), metrics.get(DEPLOY_960)
    if m480 and m960:
        a4, a9 = m480["ap_per_class_corrected"], m960["ap_per_class_corrected"]
        table(doc, "4.3", "The two deployed configurations compared",
              ["Property", "YOLO26n 480", "YOLO26n 960"],
              [["Corrected mAP@0.5 (%)", pct(m480["map50_corrected"]), pct(m960["map50_corrected"])],
               ["Precision (%)", pct(m480["precision"]), pct(m960["precision"])],
               ["Recall (%)", pct(m480["recall"]), pct(m960["recall"])],
               ["F1 (%)", pct(m480["f1"]), pct(m960["f1"])],
               ["AP closed_eye (%)", pct(a4["closed_eye"]), pct(a9["closed_eye"])],
               ["AP open_eye (%)", pct(a4["open_eye"]), pct(a9["open_eye"])],
               ["AP yawning (%)", pct(a4["yawning"]), pct(a9["yawning"])]])
    for t in C2.CH4_PERCLASS:
        body_text(doc, t)

    figure(doc, INFO / "yolo26n" / f"6-weakdevice-480-worstcase-yolo26n-test-result"
           / "tested-images" / "charts" / "01_confusion_matrix.png", "4.3",
           "Confusion matrix, deployed 480-pixel configuration")
    figure(doc, INFO / "yolo26n" / "6-weakdevice-480-worstcase-yolo26n-test-result"
           / "tested-images" / "charts" / "02_precision_recall_curve.png", "4.4",
           "Precision-recall curves, deployed 480-pixel configuration")
    figure(doc, INFO / "yolo26n" / "6-weakdevice-480-worstcase-yolo26n-test-result"
           / "training-curves" / "01_loss_curves.png", "4.5",
           "Training and validation loss, deployed 480-pixel configuration")

    heading(doc, "4.5 Measured Inference Performance", 2)
    for t in C2.CH4_LATENCY:
        body_text(doc, t)

    lrows = []
    for r in rows:
        lat = latency.get(r["key"])
        if not lat:
            continue
        pm = lat.get("parameters_millions")
        sz = lat.get("file_size_bytes")
        lrows.append([r["disp"], r["imgsz"],
                      f"{pm:.2f}" if pm else "-",
                      f"{sz / 1024 / 1024:.1f}" if sz else "-",
                      ms(lat, "cuda"), ms(lat, "cuda", "fps"),
                      ms(lat, "cpu"), ms(lat, "cpu", "fps")])
    if lrows:
        table(doc, "4.4", "Measured inference latency and model size",
              ["Model", "Input", "Params (M)", "Size (MB)",
               "GPU median (ms)", "GPU FPS", "CPU median (ms)", "CPU FPS"], lrows)
    for t in C2.CH4_LATENCY_DISCUSS:
        body_text(doc, t)

    heading(doc, "4.6 Export and Precision Conversion", 2)
    for t in C2.CH4_EXPORT:
        body_text(doc, t)

    heading(doc, "4.7 Deployed System", 2)
    for t in C2.CH4_DEPLOY:
        body_text(doc, t)
    figure(doc, INFO / "yolo26n" / "6-weakdevice-480-worstcase-yolo26n-test-result"
           / "tested-images" / "charts" / "10_sample_predictions_grid.jpg", "4.6",
           "Sample detections on held-out test images")

    heading(doc, "4.8 From-Scratch Faster R-CNN: Separate Evaluation", 2)
    for t in C2.CH4_SCRATCH:
        body_text(doc, t)
    table(doc, "4.5",
          "From-scratch Faster R-CNN results (5,705-image partition; not "
          "comparable with Table 4.2)",
          ["Metric", "Baseline", "Tuned"],
          [["mAP@0.5 (%)", "72.61", "74.27"],
           ["mAP@0.5:0.95 (%)", "32.76", "34.60"],
           ["Precision (%)", "70.62", "71.02"],
           ["Recall (%)", "82.47", "82.60"],
           ["F1 (%)", "76.09", "76.37"],
           ["AP closed_eye (%)", "74.65", "73.15"],
           ["AP open_eye (%)", "64.70", "67.57"],
           ["AP yawn (%)", "78.47", "82.10"]])
    figure(doc, SCRATCH_FS / "comparison.png", "4.7",
           "From-scratch Faster R-CNN: baseline against tuned configuration")
    figure(doc, SCRATCH_FS / "loss_curve_tuned.png", "4.8",
           "From-scratch Faster R-CNN: four-component training loss")


if __name__ == "__main__":
    sys.exit(build())
